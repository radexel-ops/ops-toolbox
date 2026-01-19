import os
import shutil
from math import ceil
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageChops
from docx import Document
from docx.shared import Cm

# UI 관련 라이브러리 (tkinter, ttk)
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

############################################################
# textsize() 대체 함수
############################################################
def my_textsize(draw_obj, text, font):
    """
    Pillow 구버전에서는 draw_obj.textsize를 지원하지만,
    특정 환경(PIL과 혼합되어 있다거나)에서 textsize가 없다는 오류가 발생할 수 있음.
    이를 대비해, textsize()가 없으면 textbbox()로 (width, height) 구함.
    """
    # textsize가 존재하면 그대로 사용
    if hasattr(draw_obj, 'textsize'):
        # (width, height) 튜플 반환
        return draw_obj.textsize(text, font=font)
    else:
        # textbbox로 대체
        if not text.strip():
            text = "M"  # 빈 문자열이면 최소 크기 보장을 위해 "M"으로 측정
        left, top, right, bottom = draw_obj.textbbox((0,0), text, font=font)
        width = right - left
        height = bottom - top
        return (width, height)


############################################################
# 헬퍼 함수: 디렉터리 내부 파일 삭제
############################################################
def clear_directory(dir_path):
    if os.path.exists(dir_path):
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f"파일 삭제 실패: {file_path}, 이유: {e}")


############################################################
# 1) QR + 텍스트 라벨 생성
############################################################
def parse_filename(filename):
    """
    파일명(확장자 제외)에서 4줄(line1, line2, line3, line4)을 추출하는 예시.
    예) "I-GW-BLG2017-0601AA.JPG" 인 경우:
      name_only = "I-GW-BLG2017-0601AA"
      - line2 = "I-GW-BLG"
      - line3 = "2017-06"
      - line4 = "01AA"
      - line1 = "" (빈 문자열)
    """
    name_only, _ = os.path.splitext(filename)
    line2 = name_only[0:8]
    line3 = name_only[8:15]
    line4 = name_only[15:]
    line1 = ""
    return line1, line2, line3, line4


def create_label(
        qr_path,
        x_path,
        output_path,
        font_path="SGr-Iosevka-Light.ttc",
        font_size=50,
        line_spacing=0,
        font_stretch_x=1.0,
        font_stretch_y=1.0,
        x_scale=0.25,
        x_size=None,
        x_offset=None,  # 기본=QR중앙
        y_offset=None,  # 기본=QR중앙
        left_margin=10,
        auto_fontsize=True,
        add_border=True,
        border_color=(0, 0, 0),
        border_width=3,
        outer_margin=20,
        text_color=(0, 0, 0),
        line_colors=None
):
    """
    2:1 구조 라벨 (왼: QR, 오른: 텍스트 4줄).
    X 이미지( x_path )를 QR 중앙(기본) 또는 임의 좌표(x_offset, y_offset)에 합성.
    X PNG에 2픽셀짜리 실제 'X' 모양의 테두리를 추가.
    """

    # ----------------------
    # 1) QR 열기
    # ----------------------
    qr_img = Image.open(qr_path).convert("RGBA")
    qr_w, qr_h = qr_img.size

    # ----------------------
    # 2) X 이미지 열고 크기 조정
    # ----------------------
    x_img = Image.open(x_path).convert("RGBA")
    if x_size is not None:
        x_new_w, x_new_h = x_size
    else:
        x_new_w = int(qr_w * x_scale)
        orig_w, orig_h = x_img.size
        ratio = orig_h / orig_w
        x_new_h = int(x_new_w * ratio)
    x_img = x_img.resize((x_new_w, x_new_h), Image.LANCZOS)

    # X PNG에 외곽선(실제 X 모양) 추가
    stroke_size = 5  # 외곽선 두께(px)
    alpha = x_img.split()[-1]  # 알파 채널
    alpha_dilated = alpha.filter(ImageFilter.MaxFilter((2 * stroke_size) + 1))
    border_mask = ImageChops.subtract(alpha_dilated, alpha)
    # 여기서 외곽선 색상을 흰색으로 지정 (255,255,255,0)
    stroke_layer = Image.new("RGBA", x_img.size, (255, 255, 255, 0))
    stroke_layer.putalpha(border_mask)
    x_with_stroke = Image.alpha_composite(stroke_layer, x_img)
    x_img = x_with_stroke

    # ----------------------
    # 3) 최종 라벨(흰 배경)
    # ----------------------
    final_w = 2 * qr_w + 3 * outer_margin
    final_h = qr_h + 2 * outer_margin
    label_img = Image.new("RGBA", (final_w, final_h), (255, 255, 255, 255))

    # ----------------------
    # 4) 라벨 왼쪽에 QR 붙이기
    # ----------------------
    label_img.paste(qr_img, (outer_margin, outer_margin))

    # ----------------------
    # 5) 오른쪽 텍스트 영역
    # ----------------------
    text_area_w = qr_w
    text_area_h = qr_h
    text_area = Image.new("RGBA", (text_area_w, text_area_h), (255, 255, 255, 255))
    text_paste_x = outer_margin + qr_w + outer_margin
    text_paste_y = outer_margin

    # ----------------------
    # 6) 파일명 -> 4줄 텍스트
    # ----------------------
    base_filename = os.path.basename(qr_path)
    line1, line2, line3, line4 = parse_filename(base_filename)
    lines = [line1, line2, line3, line4]

    # ----------------------
    # 7) 폰트 자동 축소 계산
    # ----------------------
    draw_for_calc = ImageDraw.Draw(text_area)

    def calc_text_block_size(fs, spacing):
        tmp_font = ImageFont.truetype(font_path, fs)
        total_height = 0
        max_w = 0
        for i, txt in enumerate(lines):
            measure_text = txt if txt.strip() else "M"
            w, h = my_textsize(draw_for_calc, measure_text, tmp_font)
            max_w = max(max_w, w)
            total_height += h
            if i < len(lines) - 1:
                total_height += spacing
        return tmp_font, max_w, total_height

    if auto_fontsize:
        candidate_size = font_size
        chosen_font = None
        while candidate_size > 1:
            test_font, max_w, total_h = calc_text_block_size(candidate_size, line_spacing)
            if total_h <= text_area_h - 10:
                chosen_font = test_font
                break
            candidate_size -= 1
        if chosen_font is None:
            chosen_font = ImageFont.truetype(font_path, 1)
        font = chosen_font
    else:
        font = ImageFont.truetype(font_path, font_size)

    # ----------------------
    # 8) 텍스트 그리기 -> bbox -> 리사이즈
    # ----------------------
    temp_img = Image.new("RGBA", (text_area_w, text_area_h), (255, 255, 255, 0))
    temp_draw = ImageDraw.Draw(temp_img)
    current_y = 0
    for i, txt in enumerate(lines):
        measure_text = txt if txt.strip() else "M"
        w, h = my_textsize(temp_draw, measure_text, font)
        if line_colors and i < len(line_colors):
            this_line_color = line_colors[i]
        else:
            this_line_color = text_color
        temp_draw.text((left_margin, current_y), txt, fill=this_line_color, font=font)
        current_y += h
        if i < len(lines) - 1:
            current_y += line_spacing

    bbox = temp_img.getbbox()
    if bbox:
        left, top, right, bottom = bbox
        text_crop = temp_img.crop((left, top, right, bottom))
        text_crop_w, text_crop_h = text_crop.size
    else:
        text_crop = temp_img
        text_crop_w, text_crop_h = text_area_w, text_area_h
    scaled_w = int(text_crop_w * font_stretch_x)
    scaled_h = int(text_crop_h * font_stretch_y)
    scaled_w = max(1, scaled_w)
    scaled_h = max(1, scaled_h)
    text_scaled = text_crop.resize((scaled_w, scaled_h), Image.LANCZOS)
    if scaled_w > text_area_w or scaled_h > text_area_h:
        w_ratio = text_area_w / float(scaled_w)
        h_ratio = text_area_h / float(scaled_h)
        ratio = min(w_ratio, h_ratio)
        new_w = int(scaled_w * ratio)
        new_h = int(scaled_h * ratio)
        text_scaled = text_scaled.resize((new_w, new_h), Image.LANCZOS)
        scaled_w, scaled_h = new_w, new_h
    place_x = 0
    place_y = (text_area_h - scaled_h) // 2
    text_area.paste(text_scaled, (place_x, place_y), text_scaled)
    label_img.paste(text_area, (text_paste_x, text_paste_y))

    # ----------------------
    # 9) 'X'이미지 합성 (기본: QR 코드 중앙)
    # ----------------------
    if x_offset is None:
        x_offset = outer_margin + (qr_w - x_img.width) // 2
    if y_offset is None:
        y_offset = outer_margin + (qr_h - x_img.height) // 2
    label_img.alpha_composite(x_img, (x_offset, y_offset))

    # ----------------------
    # 10) 테두리 (검은색)
    # ----------------------
    if add_border and border_width > 0:
        draw_border = ImageDraw.Draw(label_img)
        draw_border.rectangle(
            [(0, 0), (final_w - 1, final_h - 1)],
            outline=(0, 0, 0),
            width=border_width
        )

    # ----------------------
    # 11) 저장
    # ----------------------
    label_img.save(output_path)
    print(f"[완료] {output_path} 저장")


def batch_process_smplyqr():
    """
    smply-qr-codes 폴더 내 .png/.jpg/.jpeg 파일을 불러와서
    create_label() 실행 -> labels 폴더에 저장.
    실행 시 기존 labels 폴더 내부 파일은 모두 삭제.
    """
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    qr_folder = os.path.join(BASE_DIR, "smply-qr-codes")
    out_folder = os.path.join(BASE_DIR, "labels")
    os.makedirs(out_folder, exist_ok=True)
    clear_directory(out_folder)  # 기존 파일 삭제

    my_line_colors = [
        (255, 255, 255),
        (0, 0, 0),
        (0, 0, 0),
        (0, 0, 0)
    ]

    valid_ext = (".png", ".jpg", ".jpeg")
    for filename in os.listdir(qr_folder):
        if not filename.lower().endswith(valid_ext):
            continue
        qr_path = os.path.join(qr_folder, filename)
        base_name, _ = os.path.splitext(filename)
        out_name = base_name + "_label.png"
        out_path = os.path.join(out_folder, out_name)
        x_img_path = os.path.join(BASE_DIR, "X.png")
        print(f"--- 처리중: {filename} -> {out_name}")
        create_label(
            qr_path=qr_path,
            x_path=x_img_path,
            output_path=out_path,
            font_path="SGr-Iosevka-Light.ttc",
            font_size=170,
            line_spacing=10,
            font_stretch_x=1.0,
            font_stretch_y=1.0,
            x_scale=0.2,
            x_size=None,
            x_offset=None,
            y_offset=None,
            left_margin=0,
            auto_fontsize=True,
            add_border=True,
            border_color=(255, 255, 255),
            border_width=10,
            outer_margin=20,
            text_color=(0, 0, 0),
            line_colors=my_line_colors
        )


############################################################
# 2) Word 템플릿 여러 페이지 생성
############################################################
def fill_labels_in_one_doc(
    docx_path,
    image_files,
    labels_folder,
    image_width_cm=3.36,
    image_height_cm=1.69
):
    doc = Document(docx_path)
    table = doc.tables[0]
    total_cells = sum(len(row.cells) for row in table.rows)
    print(f"[INFO] 한 문서에 삽입 가능 셀: {total_cells}")
    img_idx = 0
    cell_idx = 0
    for row in table.rows:
        for cell in row.cells:
            if img_idx >= len(image_files):
                break
            img_name = image_files[img_idx]
            img_path = os.path.join(labels_folder, img_name)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(
                img_path,
                width=Cm(image_width_cm),
                height=Cm(image_height_cm)
            )
            print(f"Cell {cell_idx} ← {img_name}")
            img_idx += 1
            cell_idx += 1
        if img_idx >= len(image_files):
            break
    return doc


def fill_labels_paginated(
    template_docx="SOFT1_1164253387.docx",
    labels_folder="labels",
    output_prefix="NewLabels",
    page_capacity=60,
    image_width_cm=3.36,
    image_height_cm=1.69
):
    valid_ext = (".png", ".jpg", ".jpeg")
    image_files = sorted(
        f for f in os.listdir(labels_folder)
        if f.lower().endswith(valid_ext)
    )
    n = len(image_files)
    print(f"[INFO] 총 라벨 이미지 수: {n}")
    if n == 0:
        print("[경고] 라벨 이미지가 없습니다.")
        return
    n_docs = ceil(n / page_capacity)
    print(f"[INFO] 문서 {n_docs}개 생성 예정 (각 {page_capacity}개)")
    start_idx = 0
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    printing_folder = os.path.join(BASE_DIR, "for printing")
    os.makedirs(printing_folder, exist_ok=True)
    clear_directory(printing_folder)  # 기존 파일 삭제

    for doc_idx in range(1, n_docs + 1):
        end_idx = start_idx + page_capacity
        chunk = image_files[start_idx:end_idx]
        print(f"\n>>> [DOC {doc_idx}] 이미지 {len(chunk)}개 삽입 예정")
        doc = fill_labels_in_one_doc(
            docx_path=template_docx,
            image_files=chunk,
            labels_folder=labels_folder,
            image_width_cm=image_width_cm,
            image_height_cm=image_height_cm
        )
        out_name = f"{output_prefix}_{doc_idx}.docx"
        out_path = os.path.join(printing_folder, out_name)
        doc.save(out_path)
        print(f"[완료] {out_path} 저장 (라벨 {len(chunk)}/{n}개)")
        start_idx = end_idx


############################################################
# 3) 기본 동작 (기존 main)
############################################################
def main():
    """
    1) smply-qr-codes 폴더 안의 QR이미지 → labels 폴더 (QR+텍스트+X이미지 합성, X 테두리 적용)
    2) labels 폴더 → Word 템플릿에 60개씩 삽입 → for printing 폴더
    """
    print("=== 1) 라벨 이미지 생성 (batch_process_smplyqr) ===")
    batch_process_smplyqr()

    print("\n=== 2) 라벨 이미지를 Word 템플릿에 삽입 (fill_labels_paginated) ===")
    fill_labels_paginated(
        template_docx="SOFT1_1164253387.docx",
        labels_folder="labels",
        output_prefix="NewLabels",
        page_capacity=60,
        image_width_cm=3.36,
        image_height_cm=1.69
    )


############################################################
# 4) 사용자 지정 QR파일 리스트 처리
############################################################
def process_qr_files(qr_files):
    """
    사용자가 지정한 QR 파일 리스트(qr_files)를 받아서
    1) labels 폴더에 라벨 이미지를 생성
    2) fill_labels_paginated() 실행하여 'for printing' 폴더에 Word 문서 생성
    실행 시 기존 labels 폴더 내부 파일은 모두 삭제됨.
    """
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    out_folder = os.path.join(BASE_DIR, "labels")
    os.makedirs(out_folder, exist_ok=True)
    clear_directory(out_folder)  # 기존 파일 삭제

    my_line_colors = [
        (255, 255, 255),
        (0, 0, 0),
        (0, 0, 0),
        (0, 0, 0)
    ]

    x_img_path = os.path.join(BASE_DIR, "X.png")
    valid_ext = (".png", ".jpg", ".jpeg")

    # 1) 라벨 생성
    for filepath in qr_files:
        filename = os.path.basename(filepath)
        _, ext = os.path.splitext(filename)
        if ext.lower() not in valid_ext:
            print(f"[스킵] 이미지 파일이 아님: {filename}")
            continue
        base_name, _ = os.path.splitext(filename)
        out_name = base_name + "_label.png"
        out_path = os.path.join(out_folder, out_name)
        print(f"--- 처리중: {filename} -> {out_name}")
        create_label(
            qr_path=filepath,
            x_path=x_img_path,
            output_path=out_path,
            font_path="SGr-Iosevka-Light.ttc",
            font_size=120,
            line_spacing=10,
            font_stretch_x=1.0,
            font_stretch_y=1.0,
            x_scale=0.2,
            x_size=None,
            x_offset=None,
            y_offset=None,
            left_margin=0,
            auto_fontsize=False,
            add_border=True,
            border_color=(255, 255, 255),
            border_width=10,
            outer_margin=20,
            text_color=(0, 0, 0),
            line_colors=my_line_colors
        )

    # 2) Word 파일 생성
    fill_labels_paginated(
        template_docx="SOFT1_1164253387.docx",
        labels_folder="labels",
        output_prefix="NewLabels",
        page_capacity=60,
        image_width_cm=3.36,
        image_height_cm=1.69
    )


############################################################
# 5) GUI
############################################################
def start_ui():
    """
    스크립트를 실행했을 때 표시될 UI 함수.
    - smply-qr-codes 폴더 사용 / 파일 선택 업로드 (2가지)
    - '파일 불러오기' 버튼 -> 선택된 방식대로 파일 목록 로딩 (기존 목록 덮어쓰기)
    - '파일 추가' 버튼 -> 기존 목록에 새로운 파일들 추가
    - '선택 파일 제외' 버튼 -> 트리뷰에서 선택된 파일만 제거
    - '파일 목록 전체 리셋' 버튼 -> 트리뷰 초기화
    - '라벨지 생성 실행' 버튼 -> 선택된 파일들을 대상으로 라벨/Word 파일 생성
    """

    root = tk.Tk()
    root.title("자산관리 라벨지 생성기")
    root.geometry("600x700")

    # ttk 스타일 적용 (간단한 커스터마이즈)
    style = ttk.Style()
    style.theme_use("clam")
    style.configure("Treeview",
                    background="#FFFFFF",
                    foreground="#222222",
                    rowheight=30,
                    fieldbackground="#FFFFFF",
                    bordercolor="#EFEFEF",
                    borderwidth=1)
    style.map("Treeview",
              background=[("selected", "#E0EFFF")],
              foreground=[("selected", "#000000")])

    # 안내 라벨 프레임
    frame_top = ttk.Frame(root, padding="10 10 10 10")
    frame_top.pack(fill="x")

    instructions = (
        "1) 심플리 사이트에서 '관리번호'가 파일명인 QR코드를 다운로드 (정사각형 QR).\n"
        "2) 두 가지 방법 중 하나로 QR을 등록:\n"
        "   - 'smply-qr-codes' 폴더에 넣은 뒤 폴더 방식을 선택\n"
        "   - 파일 대화상자를 통해 업로드 (파일 선택 방식)\n"
        "3) '파일 불러오기' 버튼을 눌러 목록을 가져온 뒤,\n"
        "   필요하다면 '파일 추가' 버튼으로 추가, '선택 파일 제외'로 삭제,\n"
        "   '파일 목록 전체 리셋'으로 목록 초기화.\n"
        "4) '라벨지 생성 실행' 버튼을 누르면 'labels' 폴더에 라벨 이미지가,\n"
        "   'for printing' 폴더에 인쇄용 docx가 생성됨.\n"
        "※ Word 템플릿(SOFT1_1164253387.docx)과 X.png는 스크립트와 같은 경로에 놓아야 함."
    )

    lbl_instructions = ttk.Label(frame_top, text=instructions, anchor="w", justify="left")
    lbl_instructions.pack(fill="x")

    # 옵션 선택 프레임
    frame_options = ttk.Frame(root, padding="10 10 10 10")
    frame_options.pack(fill="x")

    mode_var = tk.StringVar(value="folder")

    rb_folder = ttk.Radiobutton(frame_options, text="smply-qr-codes 폴더 사용", variable=mode_var, value="folder")
    rb_file = ttk.Radiobutton(frame_options, text="파일 선택 업로드", variable=mode_var, value="file")

    rb_folder.grid(row=0, column=0, sticky="w", padx=5)
    rb_file.grid(row=0, column=1, sticky="w", padx=5)

    # 업로드된 파일 목록 (트리뷰)
    frame_tree = ttk.Frame(root, padding="10 10 10 10")
    frame_tree.pack(fill="both", expand=True)

    columns = ("filename", )
    file_tree = ttk.Treeview(frame_tree, columns=columns, show="headings", selectmode="extended", height=10)
    file_tree.heading("filename", text="파일명")
    file_tree.column("filename", width=600)

    vsb = ttk.Scrollbar(frame_tree, orient="vertical", command=file_tree.yview)
    file_tree.configure(yscroll=vsb.set)

    file_tree.pack(side="left", fill="both", expand=True)
    vsb.pack(side="right", fill="y")

    selected_files = []  # 업로드된 파일들의 전체 경로를 담는 리스트

    def refresh_treeview(file_list):
        for item in file_tree.get_children():
            file_tree.delete(item)
        for f in file_list:
            fname = os.path.basename(f)
            file_tree.insert("", "end", values=(fname,))

    # 파일 불러오기 (기존 목록 덮어쓰기)
    def on_load_files():
        nonlocal selected_files
        selected_files.clear()
        refresh_treeview([])

        mode = mode_var.get()
        if mode == "folder":
            base_dir = os.path.dirname(os.path.abspath(__file__))
            folder_path = os.path.join(base_dir, "smply-qr-codes")
            if not os.path.isdir(folder_path):
                messagebox.showwarning("경고", "smply-qr-codes 폴더가 존재하지 않습니다.")
                return
            valid_ext = (".png", ".jpg", ".jpeg")
            for filename in os.listdir(folder_path):
                if filename.lower().endswith(valid_ext):
                    full_path = os.path.join(folder_path, filename)
                    selected_files.append(full_path)
        else:  # mode == "file"
            filepaths = filedialog.askopenfilenames(
                title="QR 코드 이미지 선택",
                filetypes=[("Image files", "*.png;*.jpg;*.jpeg")]
            )
            selected_files = list(filepaths)

        refresh_treeview(selected_files)
        messagebox.showinfo("알림", f"총 {len(selected_files)}개 파일이 불러와졌습니다.")

    def on_add_files():
        filepaths = filedialog.askopenfilenames(
            title="추가할 QR 코드 이미지 선택",
            filetypes=[("Image files", "*.png;*.jpg;*.jpeg")]
        )
        if not filepaths:
            return
        count_added = 0
        for f in filepaths:
            if f not in selected_files:
                selected_files.append(f)
                count_added += 1
        refresh_treeview(selected_files)
        messagebox.showinfo("알림", f"새로 추가된 파일: {count_added}개")

    def on_remove_selected():
        selection = file_tree.selection()
        if not selection:
            messagebox.showwarning("경고", "제거할 파일을 선택하세요.")
            return
        removed_count = 0
        for sel in selection:
            fname = file_tree.item(sel, "values")[0]
            idx_to_remove = []
            for i, path in enumerate(selected_files):
                if os.path.basename(path) == fname:
                    idx_to_remove.append(i)
            for i in reversed(idx_to_remove):
                selected_files.pop(i)
                removed_count += 1
        refresh_treeview(selected_files)
        messagebox.showinfo("알림", f"선택된 파일 {removed_count}개가 제외되었습니다.")

    def on_reset_files():
        selected_files.clear()
        refresh_treeview([])
        messagebox.showinfo("알림", "파일 목록이 모두 초기화되었습니다.")

    def on_run():
        if not selected_files:
            messagebox.showwarning("경고", "생성할 파일 목록이 없습니다.")
            return
        process_qr_files(selected_files)
        messagebox.showinfo("완료", "'labels' 폴더와 'for printing' 폴더가 생성되었습니다.")

    frame_btns = ttk.Frame(root, padding="10 10 10 10")
    frame_btns.pack(fill="x")

    btn_load = ttk.Button(frame_btns, text="파일 불러오기", command=on_load_files)
    btn_add = ttk.Button(frame_btns, text="파일 추가", command=on_add_files)
    btn_remove = ttk.Button(frame_btns, text="선택 파일 제외", command=on_remove_selected)
    btn_reset = ttk.Button(frame_btns, text="파일 목록 전체 리셋", command=on_reset_files)

    btn_load.grid(row=0, column=0, padx=5)
    btn_add.grid(row=0, column=1, padx=5)
    btn_remove.grid(row=0, column=2, padx=5)
    btn_reset.grid(row=0, column=3, padx=5)

    btn_run = ttk.Button(root, text="라벨지 생성 실행", command=on_run)
    btn_run.pack(pady=10)

    root.mainloop()


############################################################
# 실행 진입점
############################################################
if __name__ == "__main__":
    start_ui()
